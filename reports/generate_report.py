#!/usr/bin/env python3
# Copyright 2026 The torch-spyre Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate a profiling scrum status report with charts, markdown, and Word output.

Usage:
    # Collect data from GitHub first, then generate:
    python generate_report.py

    # Or specify a custom reporting period (days):
    python generate_report.py --days 14

    # Specify a custom output directory:
    python generate_report.py --output-dir /path/to/output

The script expects JSON data files in a temp directory. If they don't exist,
it will fetch them from GitHub using the `gh` CLI.
"""

import argparse
import json
import os
import subprocess
import sys
from datetime import datetime, timedelta, timezone

import matplotlib
matplotlib.use("Agg")
import matplotlib.colors as mcolors  # noqa: E402
import matplotlib.pyplot as plt  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402

# --- Spyre brand colors ---
C_ORANGE = "#e68244"
C_GRAY = "#4d5c5e"
C_MAUVE = "#734761"
C_BLUE = "#0075ca"
C_GOLD = "#f7b749"
C_PINK = "#d9306a"
C_GREEN = "#2ea44f"

REPO = "torch-spyre/torch-spyre"


def parse_args():
    parser = argparse.ArgumentParser(description="Generate profiling scrum report")
    parser.add_argument(
        "--days", type=int, default=7, help="Reporting period in days (default: 7)"
    )
    parser.add_argument(
        "--output-dir",
        default=None,
        help="Output directory (default: directory containing this script)",
    )
    parser.add_argument(
        "--data-dir",
        default=None,
        help="Directory with cached JSON data files (default: /tmp)",
    )
    return parser.parse_args()


# --- GitHub data fetching ---
def gh_fetch(args_list, output_path):
    """Run a gh CLI command and save JSON output to a file."""
    cmd = ["gh"] + args_list
    print(f"  Fetching: {' '.join(cmd[:8])}...")
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  WARNING: gh command failed: {result.stderr.strip()}")
        with open(output_path, "w") as f:
            json.dump([], f)
        return
    with open(output_path, "w") as f:
        f.write(result.stdout)


def fetch_all_data(data_dir):
    """Fetch all profiling data from GitHub."""
    print("Fetching data from GitHub...")

    gh_fetch(
        [
            "issue", "list", "--repo", REPO, "--label", "torch-profiler",
            "--state", "open",
            "--json", "number,title,assignees,labels,updatedAt,createdAt",
            "--limit", "100",
        ],
        os.path.join(data_dir, "profiling_open_issues.json"),
    )
    gh_fetch(
        [
            "issue", "list", "--repo", REPO, "--label", "torch-profiler",
            "--state", "closed",
            "--json", "number,title,assignees,closedAt",
            "--limit", "50",
        ],
        os.path.join(data_dir, "profiling_closed_issues.json"),
    )
    gh_fetch(
        [
            "issue", "list", "--repo", REPO, "--label", "torch-profiler",
            "--label", "epic", "--state", "open",
            "--json", "number,title,assignees,body",
            "--limit", "20",
        ],
        os.path.join(data_dir, "profiling_epics.json"),
    )
    gh_fetch(
        [
            "pr", "list", "--repo", REPO, "--label", "torch-profiler",
            "--state", "open",
            "--json",
            "number,title,author,reviewRequests,reviews,createdAt,updatedAt,isDraft",
            "--limit", "30",
        ],
        os.path.join(data_dir, "profiling_prs_label.json"),
    )
    gh_fetch(
        [
            "pr", "list", "--repo", REPO, "--state", "open",
            "--search", "profil OR memory OR aiu-smi OR libaiupti",
            "--json",
            "number,title,author,reviewRequests,reviews,createdAt,isDraft,updatedAt,labels",
            "--limit", "30",
        ],
        os.path.join(data_dir, "profiling_prs_search.json"),
    )
    gh_fetch(
        [
            "pr", "list", "--repo", REPO, "--state", "merged",
            "--search", "profil OR memory OR aiu-smi OR libaiupti",
            "--json", "number,title,author,mergedAt,labels",
            "--limit", "20",
        ],
        os.path.join(data_dir, "profiling_prs_merged.json"),
    )
    print("Data fetching complete.")


# --- Helpers ---
def load_json(path):
    with open(path) as f:
        return json.load(f)


def parse_dt(s):
    if not s:
        return None
    return datetime.fromisoformat(s.replace("Z", "+00:00"))


def fmt_date(s):
    dt = parse_dt(s)
    return dt.strftime("%Y-%m-%d") if dt else ""


def assignee_names(assignees):
    return (
        ", ".join(f"@{a['login']}" for a in assignees) if assignees else "unassigned"
    )


def truncate(s, n=40):
    return s[: n - 3] + "..." if len(s) > n else s


def status_emoji(pct):
    if pct > 60:
        return ":green_circle:"
    elif pct >= 20:
        return ":yellow_circle:"
    return ":red_circle:"


PROFILING_KEYWORDS = {
    "profil", "memory", "aiu-smi", "libaiupti", "profiler", "torch-profiler",
}


def is_profiling_pr(pr):
    labels = [lbl["name"].lower() for lbl in pr.get("labels", [])]
    if "torch-profiler" in labels:
        return True
    title_lower = pr["title"].lower()
    return any(kw in title_lower for kw in PROFILING_KEYWORDS)


# --- Chart generation ---
def generate_charts(
    charts_dir, epics, opened_in_period, closed_in_period, in_progress,
    draft_prs, needs_review, approved_prs, merged_in_period,
    assignee_counts, stale_issues, start_date, today, end_date,
):
    """Generate all 5 charts as PNG files."""
    plt.style.use("seaborn-v0_8-whitegrid")

    # Chart 1: Epic Progress
    fig, ax = plt.subplots(figsize=(10, max(4, len(epics) * 0.5)))
    if epics:
        epic_labels = [truncate(f"#{e['number']} {e['title']}") for e in epics]
        epic_pcts = [e["pct"] for e in epics]
        epic_colors = [
            C_GREEN if p > 60 else C_GOLD if p >= 20 else C_PINK for p in epic_pcts
        ]
        y_pos = range(len(epics))
        ax.barh(y_pos, epic_pcts, color=epic_colors, edgecolor="white", height=0.6)
        ax.set_yticks(list(y_pos))
        ax.set_yticklabels(epic_labels, fontsize=8)
        for idx, e in enumerate(epics):
            label = f"{e['done']}/{e['total']}" if e["total"] > 0 else "no tasks"
            ax.text(max(epic_pcts[idx] + 1, 3), idx, label, va="center", fontsize=8)
        ax.set_xlim(0, 105)
        ax.set_xlabel("Completion %")
    else:
        ax.text(
            0.5, 0.5, "No epics found", ha="center", va="center",
            transform=ax.transAxes,
        )
    ax.set_title("Epic Progress", fontsize=14, fontweight="bold")
    ax.invert_yaxis()
    plt.tight_layout()
    fig.savefig(
        os.path.join(charts_dir, "epic_progress.png"), dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # Chart 2: Issue Flow
    fig, ax = plt.subplots(figsize=(6, 4))
    categories = ["Opened", "Closed", "In Progress"]
    counts = [len(opened_in_period), len(closed_in_period), len(in_progress)]
    colors = [C_BLUE, C_GREEN, C_ORANGE]
    bars = ax.bar(categories, counts, color=colors, edgecolor="white", width=0.5)
    for bar, count in zip(bars, counts):
        ax.text(
            bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2,
            str(count), ha="center", va="bottom", fontweight="bold", fontsize=12,
        )
    ax.set_ylabel("Count")
    ax.set_title(
        f"Issue Flow \u2014 {start_date.strftime('%Y-%m-%d')} to {today}",
        fontsize=13, fontweight="bold",
    )
    ax.set_ylim(0, max(counts + [1]) * 1.3)
    plt.tight_layout()
    fig.savefig(
        os.path.join(charts_dir, "issue_flow.png"), dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # Chart 3: PR Pipeline
    fig, ax = plt.subplots(figsize=(7, 3.5))
    stages = ["Draft", "Needs Review", "Approved", "Merged (this period)"]
    stage_counts = [
        len(draft_prs), len(needs_review), len(approved_prs), len(merged_in_period),
    ]
    stage_colors = [C_GRAY, C_GOLD, C_BLUE, C_GREEN]
    y_pos = range(len(stages))
    ax.barh(y_pos, stage_counts, color=stage_colors, edgecolor="white", height=0.5)
    ax.set_yticks(list(y_pos))
    ax.set_yticklabels(stages, fontsize=10)
    for idx, c in enumerate(stage_counts):
        ax.text(c + 0.1, idx, str(c), va="center", fontweight="bold", fontsize=11)
    ax.set_xlim(0, max(stage_counts + [1]) * 1.4)
    ax.set_xlabel("Count")
    ax.set_title("PR Pipeline", fontsize=13, fontweight="bold")
    ax.invert_yaxis()
    plt.tight_layout()
    fig.savefig(
        os.path.join(charts_dir, "pr_pipeline.png"), dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # Chart 4: Workload Distribution
    fig, ax = plt.subplots(figsize=(8, max(3, len(assignee_counts) * 0.4)))
    if assignee_counts:
        logins = list(assignee_counts.keys())
        issue_counts = list(assignee_counts.values())
        palette = [C_ORANGE, C_BLUE, C_MAUVE, C_GOLD, C_PINK, C_GREEN, C_GRAY]
        bar_colors = [palette[i % len(palette)] for i in range(len(logins))]
        ax.barh(logins, issue_counts, color=bar_colors, edgecolor="white", height=0.5)
        for idx, c in enumerate(issue_counts):
            ax.text(c + 0.1, idx, str(c), va="center", fontweight="bold", fontsize=10)
        ax.set_xlabel("Open Issues")
        ax.invert_yaxis()
    else:
        ax.text(
            0.5, 0.5, "No assignees", ha="center", va="center",
            transform=ax.transAxes,
        )
    ax.set_title("Open Issues by Assignee", fontsize=13, fontweight="bold")
    plt.tight_layout()
    fig.savefig(
        os.path.join(charts_dir, "workload_distribution.png"),
        dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    # Chart 5: Stale Issues Age
    fig, ax = plt.subplots(figsize=(10, max(3, len(stale_issues) * 0.45)))
    if stale_issues:
        stale_sorted = sorted(stale_issues, key=lambda i: parse_dt(i["updatedAt"]))
        stale_labels = [
            truncate(f"#{i['number']} {i['title']}", 45)
            + f" ({assignee_names(i['assignees'])})"
            for i in stale_sorted
        ]
        stale_days = [
            (end_date - parse_dt(i["updatedAt"])).days for i in stale_sorted
        ]
        norm = plt.Normalize(14, max(stale_days + [30]))
        cmap = mcolors.LinearSegmentedColormap.from_list("stale", [C_GOLD, C_PINK])
        stale_colors = [cmap(norm(d)) for d in stale_days]
        y_pos = range(len(stale_sorted))
        ax.barh(
            list(y_pos), stale_days, color=stale_colors, edgecolor="white", height=0.6,
        )
        ax.set_yticks(list(y_pos))
        ax.set_yticklabels(stale_labels, fontsize=7)
        for idx, d in enumerate(stale_days):
            ax.text(d + 0.3, idx, f"{d}d", va="center", fontsize=8)
        ax.set_xlabel("Days Since Last Update")
        ax.invert_yaxis()
    else:
        ax.text(
            0.5, 0.5, "No stale issues", ha="center", va="center",
            transform=ax.transAxes,
        )
    ax.set_title(
        "Stale Issues \u2014 Days Since Last Update", fontsize=13, fontweight="bold",
    )
    plt.tight_layout()
    fig.savefig(
        os.path.join(charts_dir, "stale_issues.png"), dpi=150, bbox_inches="tight",
    )
    plt.close(fig)

    print("Charts generated successfully.")


# --- Markdown report ---
def generate_markdown(
    output_dir, charts_dir, today, start_date, epics, closed_in_period,
    opened_in_period, in_progress, needs_review, merged_in_period, draft_prs,
    stale_issues, non_epic_open, all_open_prs, end_date, blockers,
):
    lines = []
    lines.append(f"# Profiling Scrum Status \u2014 {today}")
    lines.append("")
    lines.append(
        f"**Reporting period:** {start_date.strftime('%Y-%m-%d')} \u2192 {today}"
    )
    lines.append("")
    lines.append("---")
    lines.append("")

    # Epic Progress
    lines.append("## Epic Progress")
    lines.append("")
    lines.append("![Epic Progress](charts/epic_progress.png)")
    lines.append("")
    lines.append("| Epic | Owner | Progress | Status |")
    lines.append("|------|-------|----------|--------|")
    for e in epics:
        prog = (
            f"{e['done']}/{e['total']} ({e['pct']}%)" if e["total"] > 0 else "no tasks"
        )
        lines.append(
            f"| #{e['number']} {e['title']} | {e['assignees']}"
            f" | {prog} | {status_emoji(e['pct'])} |"
        )
    lines.append("")
    lines.append(
        "Status emoji key: :green_circle: on track (>60%),"
        " :yellow_circle: in progress (20-60%),"
        " :red_circle: blocked or behind (<20%)"
    )
    lines.append("")
    lines.append("---")
    lines.append("")

    # Issues Closed
    lines.append("## Issues Closed")
    lines.append("")
    if closed_in_period:
        lines.append("| Issue | Title | Closed by | Date |")
        lines.append("|-------|-------|-----------|------|")
        for i in sorted(
            closed_in_period, key=lambda x: x.get("closedAt", ""), reverse=True,
        ):
            lines.append(
                f"| #{i['number']} | {i['title']}"
                f" | {assignee_names(i.get('assignees', []))}"
                f" | {fmt_date(i.get('closedAt'))} |"
            )
    else:
        lines.append("None this period.")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Issues Opened
    lines.append("## Issues Opened")
    lines.append("")
    if opened_in_period:
        lines.append("| Issue | Title | Assignee | Date |")
        lines.append("|-------|-------|----------|------|")
        for i in sorted(
            opened_in_period, key=lambda x: x.get("createdAt", ""), reverse=True,
        ):
            lines.append(
                f"| #{i['number']} | {i['title']}"
                f" | {assignee_names(i.get('assignees', []))}"
                f" | {fmt_date(i.get('createdAt'))} |"
            )
    else:
        lines.append("None this period.")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Issues In Progress
    lines.append("## Issues In Progress")
    lines.append("")
    if in_progress:
        lines.append("| Issue | Title | Assignee | Last updated |")
        lines.append("|-------|-------|----------|--------------|")
        for i in sorted(
            in_progress, key=lambda x: x.get("updatedAt", ""), reverse=True,
        ):
            lines.append(
                f"| #{i['number']} | {i['title']}"
                f" | {assignee_names(i.get('assignees', []))}"
                f" | {fmt_date(i.get('updatedAt'))} |"
            )
    else:
        lines.append("None this period.")
    lines.append("")
    lines.append("---")
    lines.append("")

    # PRs Needing Review
    lines.append("## PRs Needing Review")
    lines.append("")
    if needs_review:
        lines.append("| PR | Title | Author | Waiting since | Reviewers |")
        lines.append("|----|-------|--------|---------------|-----------|")
        for pr in sorted(needs_review, key=lambda x: x.get("createdAt", "")):
            author = pr["author"]["login"]
            rr = ", ".join(
                f"@{r.get('login', '')}" for r in pr.get("reviewRequests", [])
            )
            days_waiting = (end_date - parse_dt(pr["createdAt"])).days
            flag = " (stale)" if days_waiting > 3 else ""
            lines.append(
                f"| #{pr['number']} | {pr['title']} | @{author}"
                f" | {fmt_date(pr['createdAt'])}{flag} | {rr} |"
            )
    else:
        lines.append("None this period.")
    lines.append("")
    lines.append("---")
    lines.append("")

    # PRs Merged
    lines.append("## PRs Merged")
    lines.append("")
    if merged_in_period:
        lines.append("| PR | Title | Author | Merged |")
        lines.append("|----|-------|--------|--------|")
        for pr in sorted(
            merged_in_period, key=lambda x: x.get("mergedAt", ""), reverse=True,
        ):
            lines.append(
                f"| #{pr['number']} | {pr['title']}"
                f" | @{pr['author']['login']} | {fmt_date(pr.get('mergedAt'))} |"
            )
    else:
        lines.append("None this period.")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Draft PRs
    lines.append("## Draft PRs")
    lines.append("")
    if draft_prs:
        lines.append("| PR | Title | Author | Created |")
        lines.append("|----|-------|--------|---------|")
        for pr in draft_prs:
            lines.append(
                f"| #{pr['number']} | {pr['title']}"
                f" | @{pr['author']['login']} | {fmt_date(pr['createdAt'])} |"
            )
    else:
        lines.append("None this period.")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Blockers & Risks
    lines.append("## Blockers & Risks")
    lines.append("")
    for b in blockers:
        lines.append(b)
    lines.append("")
    lines.append("---")
    lines.append("")

    # Key Numbers
    lines.append("## Key Numbers")
    lines.append("")
    lines.append(f"- **Open issues:** {len(non_epic_open)} (+ {len(epics)} epics)")
    lines.append(f"- **Closed this period:** {len(closed_in_period)}")
    lines.append(f"- **Open PRs (profiling):** {len(all_open_prs)}")
    lines.append(f"- **PRs needing review:** {len(needs_review)}")
    lines.append(f"- **PRs merged this period:** {len(merged_in_period)}")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Visualizations
    lines.append("## Visualizations")
    lines.append("")
    for name, fname in [
        ("Issue Flow", "issue_flow.png"),
        ("PR Pipeline", "pr_pipeline.png"),
        ("Workload Distribution", "workload_distribution.png"),
        ("Stale Issues", "stale_issues.png"),
    ]:
        lines.append(f"### {name}")
        lines.append(f"![{name}](charts/{fname})")
        lines.append("")

    md_path = os.path.join(output_dir, f"scrum-profiling-{today}.md")
    with open(md_path, "w") as f:
        f.write("\n".join(lines))
    print(f"Markdown report saved: {md_path}")
    return md_path


# --- Word document ---
def generate_docx(
    output_dir, charts_dir, today, start_date, epics, closed_in_period,
    opened_in_period, in_progress, needs_review, merged_in_period, draft_prs,
    blockers, non_epic_open, all_open_prs, end_date,
):
    doc = Document()
    doc.add_heading(f"Profiling Scrum Status \u2014 {today}", level=0)
    subtitle = doc.add_paragraph(
        f"Reporting period: {start_date.strftime('%Y-%m-%d')} to {today}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_table(headers, rows):
        table = doc.add_table(
            rows=1 + len(rows), cols=len(headers), style="Light List Accent 1",
        )
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = table.rows[i + 1].cells[j]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)

    def add_chart(name):
        path = os.path.join(charts_dir, name)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(6))

    # Epic Progress
    doc.add_heading("Epic Progress", level=2)
    add_chart("epic_progress.png")
    epic_rows = []
    for e in epics:
        prog = (
            f"{e['done']}/{e['total']} ({e['pct']}%)" if e["total"] > 0 else "no tasks"
        )
        st = (
            "on track" if e["pct"] > 60
            else "in progress" if e["pct"] >= 20
            else "behind"
        )
        epic_rows.append([f"#{e['number']} {e['title']}", e["assignees"], prog, st])
    add_table(["Epic", "Owner", "Progress", "Status"], epic_rows)

    # Issues Closed
    doc.add_heading("Issues Closed", level=2)
    if closed_in_period:
        rows = [
            [
                f"#{i['number']}", i["title"],
                assignee_names(i.get("assignees", [])), fmt_date(i.get("closedAt")),
            ]
            for i in sorted(
                closed_in_period, key=lambda x: x.get("closedAt", ""), reverse=True,
            )
        ]
        add_table(["Issue", "Title", "Closed by", "Date"], rows)
    else:
        doc.add_paragraph("None this period.")

    # Issues Opened
    doc.add_heading("Issues Opened", level=2)
    if opened_in_period:
        rows = [
            [
                f"#{i['number']}", i["title"],
                assignee_names(i.get("assignees", [])), fmt_date(i.get("createdAt")),
            ]
            for i in sorted(
                opened_in_period, key=lambda x: x.get("createdAt", ""), reverse=True,
            )
        ]
        add_table(["Issue", "Title", "Assignee", "Date"], rows)
    else:
        doc.add_paragraph("None this period.")

    # Issues In Progress
    doc.add_heading("Issues In Progress", level=2)
    if in_progress:
        rows = [
            [
                f"#{i['number']}", i["title"],
                assignee_names(i.get("assignees", [])), fmt_date(i.get("updatedAt")),
            ]
            for i in sorted(
                in_progress, key=lambda x: x.get("updatedAt", ""), reverse=True,
            )
        ]
        add_table(["Issue", "Title", "Assignee", "Last Updated"], rows)
    else:
        doc.add_paragraph("None this period.")

    # PRs Needing Review
    doc.add_heading("PRs Needing Review", level=2)
    if needs_review:
        rows = []
        for pr in sorted(needs_review, key=lambda x: x.get("createdAt", "")):
            rr = ", ".join(
                f"@{r.get('login', '')}" for r in pr.get("reviewRequests", [])
            )
            rows.append([
                f"#{pr['number']}", pr["title"],
                f"@{pr['author']['login']}", fmt_date(pr["createdAt"]), rr,
            ])
        add_table(["PR", "Title", "Author", "Waiting Since", "Reviewers"], rows)
    else:
        doc.add_paragraph("None this period.")

    # PRs Merged
    doc.add_heading("PRs Merged", level=2)
    if merged_in_period:
        rows = [
            [
                f"#{pr['number']}", pr["title"],
                f"@{pr['author']['login']}", fmt_date(pr.get("mergedAt")),
            ]
            for pr in sorted(
                merged_in_period, key=lambda x: x.get("mergedAt", ""), reverse=True,
            )
        ]
        add_table(["PR", "Title", "Author", "Merged"], rows)
    else:
        doc.add_paragraph("None this period.")

    # Draft PRs
    doc.add_heading("Draft PRs", level=2)
    if draft_prs:
        rows = [
            [
                f"#{pr['number']}", pr["title"],
                f"@{pr['author']['login']}", fmt_date(pr["createdAt"]),
            ]
            for pr in draft_prs
        ]
        add_table(["PR", "Title", "Author", "Created"], rows)
    else:
        doc.add_paragraph("None this period.")

    # Blockers & Risks
    doc.add_heading("Blockers & Risks", level=2)
    add_chart("stale_issues.png")
    for b in blockers:
        clean = b.lstrip("- ").lstrip(" ")
        doc.add_paragraph(clean, style="List Bullet")

    # Key Numbers
    doc.add_heading("Key Numbers", level=2)
    doc.add_paragraph(
        f"Open issues: {len(non_epic_open)} (+ {len(epics)} epics)",
        style="List Bullet",
    )
    doc.add_paragraph(
        f"Closed this period: {len(closed_in_period)}", style="List Bullet",
    )
    doc.add_paragraph(
        f"Open PRs (profiling): {len(all_open_prs)}", style="List Bullet",
    )
    doc.add_paragraph(
        f"PRs needing review: {len(needs_review)}", style="List Bullet",
    )
    doc.add_paragraph(
        f"PRs merged this period: {len(merged_in_period)}", style="List Bullet",
    )

    # Visualization appendix
    doc.add_heading("Visualizations", level=2)
    for chart_name, chart_title in [
        ("issue_flow.png", "Issue Flow"),
        ("pr_pipeline.png", "PR Pipeline"),
        ("workload_distribution.png", "Workload Distribution"),
    ]:
        doc.add_heading(chart_title, level=3)
        add_chart(chart_name)

    docx_path = os.path.join(output_dir, f"scrum-profiling-{today}.docx")
    doc.save(docx_path)
    print(f"Word document saved: {docx_path}")
    return docx_path


# --- Main ---
def main():
    args = parse_args()

    end_date = datetime.now(timezone.utc)
    start_date = end_date - timedelta(days=args.days)
    today = end_date.strftime("%Y-%m-%d")

    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = args.output_dir or script_dir
    data_dir = args.data_dir or "/tmp"
    charts_dir = os.path.join(output_dir, "charts")

    os.makedirs(charts_dir, exist_ok=True)

    def in_period(dt_str):
        dt = parse_dt(dt_str)
        return dt and dt >= start_date

    # Fetch data if not cached
    data_file = os.path.join(data_dir, "profiling_open_issues.json")
    if not os.path.exists(data_file):
        fetch_all_data(data_dir)
    else:
        print(f"Using cached data from {data_dir}")

    # Load data
    open_issues = load_json(os.path.join(data_dir, "profiling_open_issues.json"))
    closed_issues = load_json(os.path.join(data_dir, "profiling_closed_issues.json"))
    epics_raw = load_json(os.path.join(data_dir, "profiling_epics.json"))
    prs_label = load_json(os.path.join(data_dir, "profiling_prs_label.json"))
    prs_search = load_json(os.path.join(data_dir, "profiling_prs_search.json"))
    prs_merged_raw = load_json(os.path.join(data_dir, "profiling_prs_merged.json"))

    # Classify
    epic_numbers = {e["number"] for e in epics_raw}
    non_epic_open = [i for i in open_issues if i["number"] not in epic_numbers]

    closed_in_period = [i for i in closed_issues if in_period(i.get("closedAt"))]
    opened_in_period = [
        i for i in non_epic_open if in_period(i.get("createdAt"))
    ]
    opened_numbers = {i["number"] for i in opened_in_period}
    in_progress = [
        i for i in non_epic_open
        if i.get("assignees") and in_period(i.get("updatedAt"))
        and i["number"] not in opened_numbers
    ]

    stale_issues = [
        i for i in non_epic_open
        if i.get("assignees")
        and parse_dt(i.get("updatedAt"))
        and (end_date - parse_dt(i["updatedAt"])).days >= 14
    ]

    # Epics
    epics = []
    for e in epics_raw:
        body = e.get("body", "") or ""
        done = body.count("- [x]") + body.count("- [X]")
        total = done + body.count("- [ ]")
        pct = int(done / total * 100) if total > 0 else 0
        epics.append({
            "number": e["number"],
            "title": e["title"],
            "assignees": assignee_names(e.get("assignees", [])),
            "done": done,
            "total": total,
            "pct": pct,
        })

    # PRs
    seen_pr = set()
    all_open_prs = []
    for pr in prs_label + prs_search:
        if pr["number"] not in seen_pr:
            seen_pr.add(pr["number"])
            all_open_prs.append(pr)
    all_open_prs = [pr for pr in all_open_prs if is_profiling_pr(pr)]

    draft_prs = [pr for pr in all_open_prs if pr.get("isDraft")]
    non_draft_prs = [pr for pr in all_open_prs if not pr.get("isDraft")]

    needs_review = []
    approved_prs = []
    for pr in non_draft_prs:
        reviews = pr.get("reviews", [])
        states = [r["state"] for r in reviews]
        if "APPROVED" in states and states[-1] == "APPROVED":
            approved_prs.append(pr)
        else:
            needs_review.append(pr)

    merged_in_period = [
        pr for pr in prs_merged_raw
        if in_period(pr.get("mergedAt")) and is_profiling_pr(pr)
    ]

    # Workload
    assignee_counts = {}
    for i in non_epic_open:
        for a in i.get("assignees", []):
            login = a["login"]
            assignee_counts[login] = assignee_counts.get(login, 0) + 1
    assignee_counts = dict(sorted(assignee_counts.items(), key=lambda x: -x[1]))

    # Blockers
    blockers = []
    if stale_issues:
        blockers.append(
            f"- **{len(stale_issues)} stale issues**"
            " (assigned, no update in 14+ days):"
        )
        for i in stale_issues[:5]:
            days = (end_date - parse_dt(i["updatedAt"])).days
            blockers.append(
                f"  - #{i['number']} {truncate(i['title'], 50)}"
                f" ({assignee_names(i.get('assignees', []))}) \u2014 {days}d stale"
            )
        if len(stale_issues) > 5:
            blockers.append(f"  - ...and {len(stale_issues) - 5} more")

    unassigned_epics = [e for e in epics if e["assignees"] == "unassigned"]
    if unassigned_epics:
        blockers.append(
            f"- **{len(unassigned_epics)} unassigned epics** need owners: "
            + ", ".join(f"#{e['number']}" for e in unassigned_epics)
        )

    no_task_epics = [e for e in epics if e["total"] == 0]
    if no_task_epics:
        blockers.append(
            f"- **{len(no_task_epics)} epics have no sub-tasks**"
            " defined \u2014 hard to track progress"
        )

    if not merged_in_period:
        blockers.append(
            "- **No profiling PRs merged this period** \u2014 velocity concern"
        )

    if not blockers:
        blockers.append("None identified this period.")

    # Generate
    generate_charts(
        charts_dir, epics, opened_in_period, closed_in_period, in_progress,
        draft_prs, needs_review, approved_prs, merged_in_period,
        assignee_counts, stale_issues, start_date, today, end_date,
    )

    generate_markdown(
        output_dir, charts_dir, today, start_date, epics, closed_in_period,
        opened_in_period, in_progress, needs_review, merged_in_period, draft_prs,
        stale_issues, non_epic_open, all_open_prs, end_date, blockers,
    )

    generate_docx(
        output_dir, charts_dir, today, start_date, epics, closed_in_period,
        opened_in_period, in_progress, needs_review, merged_in_period, draft_prs,
        blockers, non_epic_open, all_open_prs, end_date,
    )

    # Summary
    print(f"\n{'=' * 40}")
    print(f"Report generated for {today}")
    print(f"{'=' * 40}")
    print(f"Open issues: {len(non_epic_open)} (+ {len(epics)} epics)")
    print(f"Closed this period: {len(closed_in_period)}")
    print(f"Open PRs (profiling): {len(all_open_prs)}")
    print(f"PRs needing review: {len(needs_review)}")
    print(f"PRs merged this period: {len(merged_in_period)}")
    print(f"Stale issues: {len(stale_issues)}")


if __name__ == "__main__":
    main()
